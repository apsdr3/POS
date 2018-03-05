#NEED TO pip install pyexcel
#NEED TO pip install pyexcel-xls
#NEED TO pip install XlsxWriter    allows user to create excel files
#NEED TO pip install openpyxl   allows user to modify excel files
#NEED TO pip install python-docx    allows user to create and edit word documents
import tkinter as tk
import pyexcel as pe
import xlsxwriter as xw
import datetime
import time

from tkinter import ttk
from openpyxl import load_workbook,Workbook
from tkinter import filedialog
from tkinter import *
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

LARGE_FONT = ("Verdana", 12)
NORMAL_FONT = ("Verdana", 10)
SMALL_FONT = ("Verdana", 8)

#GLOBAL VARIABLES
masterList = [] #list of objects from Master File
#customerList = []
customerList = [[20297939, 'AC EDT 75ML SPRAY TEST', 250, 0, 3605520297939, 2, 1], [72728313, 'ADG HOMME DEMO.TOIL.WAT.SPRAY 100ML', 250, 0, 3360372728313, 1, 1], [20267093, 'ACF EDP 75ML SPRAY TEST', 300, 0, 3605520267093, 0, 1], [20297878, 'AM EDT 100ML SPRAY TEST', 350, 0, 3605520297878, 0, 1], [70638152, 'ADGA JASMINE V100ML TEST OS', 400, 0, 3614270638152, 0, 1], [71381477, 'SUN DI GIOIA EDP V30ML', 700, 0, 3614271381477, 0, 1], [71381538, 'SUN DI GIOIA EDP V50ML TEST', 1000, 0, 3614271381538, 0, 1], [71381392, 'AIR DI GIOIA EDP V50ML', 1000, 0, 3614271381392, 0, 1], [71381491, 'SUN DI GIOIA EDP V100ML', 1300, 0, 3614271381491, 0, 1], [70157639, 'ADGH PROFUMO SP75ML', 1350, 0, 3614270157639, 0, 1], [22035423, 'SI EDT SP100ML', 1500, 0, 3605522035423, 0, 1], [71214799, 'SI EDP ROSE V100ML OS', 2200, 0, 3614271214799, 0, 1], [21583413, 'AP FIGUE EDEN EDT SP100ML', 2500, 0, 3605521583413, 0, 1], [72009436, 'DKN EDT 100ML SPRAY', 700, 0, 3360372009436, 0, 1], [30655201, 'FLUDLST SCELLG FLUID SPRAY RET 150ML', 200, 0, 3474630655201, 0, 1], [30543003, 'STY FORME FATALE 125ML', 200, 0, 3474630543003, 0, 1], [30542709, 'STY LAQUE COUTURE 300ML', 200, 0, 3474630542709, 0, 1], [36382378, 'NUT MAGISTRAL SOIN N2 CONCENT 500ML', 250, 0, 3474636382378, 0, 1], [36382361, 'NUT MAGISTRAL SOIN N1 ANCREUR 500ML', 250, 0, 3474636382361, 0, 1]]
container = 0
excelString = " "
filename = " "
paymentString = ""

#sets customer info
customerName = " "
phone = 0
address = " "
customerType = "New"
totalCustomerPayment = 0
tinValue = "N/A"
BStyleValue = "N/A"
termsValue = "N/A"
PWDValue = "N/A"
cashierString = ""

errorCode = 0
"""
Error Code Legend:
0 = No error
1 = Master File Error
"""

modeCode = 1
"""
modeCode Legend:
0 = Inventory Mode
1 = Transaction Mode
"""


class POS(tk.Tk):
    #initializes on startup
    def __init__(self,*args,**kwargs):
        tk.Tk.__init__(self, *args, **kwargs)   #initializes tk module

        #tk.Tk.iconbitmap(self, default="<image-file-name.ico>") #gives an icon for the program, top left corner, has to be an ICON
        tk.Tk.wm_title(self, "FONZY")   #Gives name to program client application
        
        global container
        container = tk.Frame(self)
        container.pack(side = "top", fill = "both", expand = True)  #"packs" or pushes the container to the top

        container.grid_rowconfigure(0, weight = 1)
        container.grid_columnconfigure(0, weight = 1)

        #sets customer info
        global customerName
        global phone
        global address

        customerName = StringVar()
        customerName.set("")
        
        phone = IntVar()
        phone.set(0)
        
        address = StringVar()
        address.set("")


        self.frames = {}    #creates an object to hold multiple frames i.e. more windows/tabs

        for F in (ErrorPage, MainPage): #Adds frames onto list, to add more frames, just add it to the list
            frame = F(container, self)
            self.frames[F] = frame  #adds frame into frames object
            frame.grid(row=0, column = 0, sticky = "nsew")  #sets frame structure, nsew = north south east west
        
        self.show_frame(MainPage)

    #shows the frame when called
    def show_frame(self,cont):
        frame = self.frames[cont]
        frame.tkraise() 




#MAIN FRAME THAT THE USER WILL BE ON! Showcases customer information and purchase details
class MainPage(tk.Frame):
    def __init__(self, parent, controller):
        tk.Frame.__init__(self,parent)
        global customerName
        global phone
        global address

        if modeCode == 1: #Transaction mode
    # FRAME1 ------------------ CUSTOMER NAME, TELEPHONE NUMBER ----------------------#

            #First frame inside the current MainPage Window Frame
            frame1 = Frame(self, bg = "green")
            frame1.pack(fill = BOTH)
                   
    #CUSTOMER NAME
            frame1Label = tk.Label(frame1, text = "Customer Name", font = SMALL_FONT)
            frame1Label.grid(row = 0, column = 0, padx = 5, pady = 5, sticky = W)

            frame1EntryBox1 = ttk.Entry(frame1, textvariable = customerName, width = 40)   #creates an entry box and allows the entry of a string variable
            frame1EntryBox1.grid(row = 0, column = 1, padx = 5, pady = 5, sticky = W)

    #PHONE
            frame1Label2 = tk.Label(frame1, text="Phone", font=SMALL_FONT)
            frame1Label2.grid(row = 0, column = 2, padx = 5, pady = 5, sticky = W)
            
            frame1EntryBox2 = ttk.Entry(frame1, textvariable = phone, width = 40)   #creates an entry box and allows the entry of a string variable
            frame1EntryBox2.grid(row = 0, column = 3, padx = 5, pady = 5, sticky = W)
    #-----------------------------------------------------------------------------#


    # FRAME2 -------------------------- ADDRESS ------------------------------------#
            
            #Second frame inside the current MainPage Window Frame
            frame2 = Frame(self, bg = "red")
            frame2.pack(fill = BOTH)

    #ADDRESS
            frame2.label = tk.Label(frame2, text="Address", font=SMALL_FONT)
            frame2.label.grid(row = 1, column = 0, padx=5, pady=5, sticky = W)
            
            frame2EntryBox = ttk.Entry(frame2, textvariable = address, width = 99)   #creates an entry box and allows the entry of a string variable
            frame2EntryBox.grid(row = 1, column = 1, padx = 5, pady = 5, sticky = W)
    #-----------------------------------------------------------------------------------#


    # FRAME3 -------------------------- ENTRY BOX GRID ------------------------------------#
            #Third frame inside the current MainPage Window Frame

            frame3 = Frame(self, bg = "white", width = 690, height = 400, borderwidth = 1)
            frame3.pack(expand = True, fill = Y)

            #Creates a canvas-frame scrollable widget
            frame3Canvas = tk.Canvas(frame3, width = 690, height = 400, borderwidth = 0, bg="white")
            frame3Frame = tk.Frame(frame3Canvas, bg = "white")
            frame3ScrollBar = tk.Scrollbar(frame3, orient = "vertical", command = frame3Canvas.yview)
            frame3Canvas.configure(yscrollcommand = frame3ScrollBar.set)

            frame3ScrollBar.pack(side = "right", fill = "y")
            frame3Canvas.pack(side = "left", fill = "both", expand = True)
            frame3Canvas.create_window((4,4), window = frame3Frame, anchor = "nw")

            frame3Frame.bind("<Configure>", lambda event, canvas = frame3Canvas: frame3Canvas.configure(scrollregion = frame3Canvas.bbox("all")))


            #THIS CREATES THE GRID TO OUTPUT THE DATA QUERIED FROM THE MASTER FILE
            #Description of Box: 'X' amount goind down, i.e. number of items, 6 descriptive columns: Bar Code, Product Description, Amount, Quantity, Additional Discount, Total Amount
            #This is the header GRID VIEW Labels
            frame3Label1 = tk.Label(frame3Frame, text = "Bar Code", font = NORMAL_FONT, relief = SUNKEN, width = 15)
            frame3Label1.grid(row = 0, column = 0)

            frame3Label2 = tk.Label(frame3Frame, text = "Product Description", font = NORMAL_FONT, relief = SUNKEN, width = 30)
            frame3Label2.grid(row = 0, column = 1)

            frame3Label3 = tk.Label(frame3Frame, text = "Price", font = NORMAL_FONT, relief = SUNKEN, width = 9)
            frame3Label3.grid(row = 0, column = 2)

            frame3Label4 = tk.Label(frame3Frame, text = "Quantity", font = NORMAL_FONT, relief = SUNKEN, width = 8)
            frame3Label4.grid(row = 0, column = 3)

            frame3Label5 = tk.Label(frame3Frame, text = "Discount", font = NORMAL_FONT, relief = SUNKEN, width = 8)
            frame3Label5.grid(row = 0, column = 4)

            frame3Label6 = tk.Label(frame3Frame, text = "Cost", font = NORMAL_FONT, relief = SUNKEN, width = 10)
            frame3Label6.grid(row = 0, column = 5)

            
            #Need to create a "Dynamically allocated grid view entry boxes" for next boxes with scroll wheel
            rowNum = 1
            totalCost = 0
            totalQuantity = 0
            for i in range(len(customerList)):

                barCodeString = str(customerList[i][4])
                frame3BarCode = tk.Label(frame3Frame, text = barCodeString, font = NORMAL_FONT, relief = SUNKEN, width = 15)
                frame3BarCode.grid(row = rowNum, column = 0)

                prodDesc = customerList[i][1]
                frame3ProdDesc = tk.Label(frame3Frame, text = prodDesc, font = NORMAL_FONT, relief = SUNKEN, width = 30)
                frame3ProdDesc.grid(row = rowNum, column = 1)

                frame3Price = tk.Label(frame3Frame, text = "{:,}".format(customerList[i][2]), font = NORMAL_FONT, relief = SUNKEN, width = 9)
                frame3Price.grid(row = rowNum, column = 2)

                totalQuantity += customerList[i][6]
                frame3Quantity = ttk.Label(frame3Frame, text = "{:,}".format(customerList[i][6]), font = NORMAL_FONT, relief = SUNKEN, width = 8)   #creates an entry box and allows the entry of a string variable
                frame3Quantity.grid(row = rowNum, column = 3)

                discountString = str(customerList[i][3])
                frame3Discount = tk.Label(frame3Frame, text = discountString+"%", font = NORMAL_FONT, relief = SUNKEN, width = 8)
                frame3Discount.grid(row = rowNum, column = 4)

                cost = (customerList[i][6]*customerList[i][2])-((customerList[i][3]/100)*(customerList[i][6]*customerList[i][2]))   #gets cost estimate with given mathematical values
                totalCost += cost
                frame3Cost = tk.Label(frame3Frame, text = "{:,}".format(cost), font = NORMAL_FONT, relief = SUNKEN, width = 10)
                frame3Cost.grid(row = rowNum, column = 5)

                rowNum += 1        
    #--------------------------------------------------------------------------------------#


    # FRAME4 -------------------------- TOTAL QUANTITY AND AMOUNT ------------------------------------#
            #Fourth frame inside the current MainPage Window Frame
            frame4 = Frame(self, bg = "blue")
            frame4.pack(fill = BOTH)

            frame4Label1 = tk.Label(frame4, text="Total Quantity", font = NORMAL_FONT)
            frame4Label1.grid(row = 0, column = 0, padx = 5, pady = 5)


            frame4Label2 = tk.Label(frame4, text = "{:,}".format(totalQuantity), font = NORMAL_FONT, relief = SUNKEN, width = 16)  #NEED TO GET SUM OF TOTAL QUANTITY
            frame4Label2.grid(row = 0, column = 1)


            #Frame block to separate the first set of labels above from the second set of labels below
            frame4LabelBlock = tk.Label(frame4, width = 16, bg = "blue")
            frame4LabelBlock.grid(row = 0, column = 2, padx = 45, pady = 5)


            frame4Label3 = tk.Label(frame4, text = "Total Cost", font=NORMAL_FONT)
            frame4Label3.grid(row = 0, column = 3, padx = 5, pady = 5)

            frame4Label4 = tk.Label(frame4, text = "{:,}".format(totalCost), font = NORMAL_FONT, relief = SUNKEN, width = 16)  #NEED TO GET SUM OF TOTAL COST
            frame4Label4.grid(row = 0, column = 4)
    #-------------------------------------------------------------------------------------------------#


    # FRAME5 -------------------------- BAR CODE AND ADD ITEM ------------------------------------#
            #Fifth frame inside the current MainPage Window Frame
            frame5 = Frame(self, bg = "pink")
            frame5.pack(fill = BOTH)

            frame5Label1 = tk.Label(frame5, text = "Product Bar Code", font = NORMAL_FONT)
            frame5Label1.grid(row = 0, column = 0, padx=5, pady=5, sticky = W)

            barCode = tk.StringVar()    #creates the object barCode with a string variable type
            frame5EntryBox = ttk.Entry(frame5, textvariable = barCode, width = 22)   #creates an entry box and allows the entry of a string variable
            frame5EntryBox.focus_set()  #cursor default on entry box
            frame5EntryBox.grid(row = 0, column = 1, padx = 5, pady = 5)

            frame5Spacer = tk.Label(frame5, text = "", font = NORMAL_FONT, bg = "pink")
            frame5Spacer.grid(row = 0, column = 2, padx=20, pady=5, sticky = W)

            frame5Label2 = tk.Label(frame5, text = "Quantity", font = NORMAL_FONT)
            frame5Label2.grid(row = 0, column = 3, padx=5, pady=5, sticky = W)

            quantity = tk.IntVar()    #creates the object barCode with a string variable type
            quantity.set(1)
            frame5EntryBox2 = ttk.Entry(frame5, textvariable = quantity, width = 8)   #creates an entry box and allows the entry of a string variable
            frame5EntryBox2.grid(row = 0, column = 4, padx = 5, pady = 5)

            frame5Button = ttk.Button(frame5, text = "Add Item", command = lambda: updateCustomerList(barCode, quantity))
            frame5Button.grid(row = 0, column = 5, padx = 90, pady = 10)

            self.winfo_toplevel().bind("<Return>", lambda event: updateCustomerList(barCode=barCode, quantity=quantity))    #binds enter/return key to add the barcode given in the entry box
    #---------------------------------------------------------------------------------------------#


    # FRAME6 -------------------------- PROCESS AND REFRESH ------------------------------------#
            #Sixth frame inside the current MainPage Window Frame
            frame6 = Frame(self, bg = "yellow")
            frame6.pack(fill = BOTH)

            frame6Button = ttk.Button(frame6, text = "Refresh", command=refreshMainFrame)
            frame6Button.grid(row = 0, column = 0, padx = 125, pady = 10)

            frame6Button = ttk.Button(frame6, text = "Process", command=beforeProcessPagePopup)
            frame6Button.grid(row = 0, column = 1, padx = 125, pady = 10)
    #--------------------------------------------------------------------------------------------#   

        else: #Inventory mode
    # FRAME1 ------------------ CUSTOMER NAME, TELEPHONE NUMBER ----------------------#

            #First frame inside the current MainPage Window Frame
            frame1 = Frame(self, bg = "green")
            frame1.pack(fill = BOTH)
            
    #CUSTOMER NAME
            frame1Label = tk.Label(frame1, text = "Customer Name", font = SMALL_FONT)
            frame1Label.grid(row = 0, column = 0, padx = 5, pady = 5, sticky = W)

            customerName = StringVar()    #creates the object customerName with a string variable type
            frame1EntryBox1 = ttk.Entry(frame1, textvariable = customerName, width = 40, state = tk.DISABLED)   #creates an entry box and allows the entry of a string variable
            frame1EntryBox1.grid(row = 0, column = 1, padx = 5, pady = 5, sticky = W)

    #PHONE
            frame1Label2 = tk.Label(frame1, text="Phone", font=SMALL_FONT)
            frame1Label2.grid(row = 0, column = 2, padx = 5, pady = 5, sticky = W)
            
            phone = StringVar()    #creates the object phone with a string variable type
            frame1EntryBox2 = ttk.Entry(frame1, textvariable = phone, width = 40, state = tk.DISABLED)   #creates an entry box and allows the entry of a string variable
            frame1EntryBox2.grid(row = 0, column = 3, padx = 5, pady = 5, sticky = W)
    #-----------------------------------------------------------------------------#


    # FRAME2 -------------------------- ADDRESS ------------------------------------#
            
            #Second frame inside the current MainPage Window Frame
            frame2 = Frame(self, bg = "red")
            frame2.pack(fill = BOTH)

    #ADDRESS
            frame2.label = tk.Label(frame2, text="Address", font=SMALL_FONT)
            frame2.label.grid(row = 1, column = 0, padx=5, pady=5, sticky = W)

            address = StringVar()    #creates the object address with a string variable type
            frame2EntryBox = ttk.Entry(frame2, textvariable = address, width = 99, state = tk.DISABLED)   #creates an entry box and allows the entry of a string variable
            frame2EntryBox.grid(row = 1, column = 1, padx = 5, pady = 5, sticky = W)
    #-----------------------------------------------------------------------------------#


    # FRAME3 -------------------------- ENTRY BOX GRID ------------------------------------#
            #Third frame inside the current MainPage Window Frame

            frame3 = Frame(self, bg = "white", width = 690, height = 400, borderwidth = 1)
            frame3.pack(expand = True, fill = Y)

            #Creates a canvas-frame scrollable widget
            frame3Canvas = tk.Canvas(frame3, width = 690, height = 400, borderwidth = 0, bg="white")
            frame3Frame = tk.Frame(frame3Canvas, bg = "white")
            frame3ScrollBar = tk.Scrollbar(frame3, orient = "vertical", command = frame3Canvas.yview)
            frame3Canvas.configure(yscrollcommand = frame3ScrollBar.set)

            frame3ScrollBar.pack(side = "right", fill = "y")
            frame3Canvas.pack(side = "left", fill = "both", expand = True)
            frame3Canvas.create_window((4,4), window = frame3Frame, anchor = "nw")

            frame3Frame.bind("<Configure>", lambda event, canvas = frame3Canvas: frame3Canvas.configure(scrollregion = frame3Canvas.bbox("all")))


            #THIS CREATES THE GRID TO OUTPUT THE DATA QUERIED FROM THE MASTER FILE
            #Description of Box: 'X' amount goind down, i.e. number of items, 6 descriptive columns: Bar Code, Product Description, Amount, Quantity, Additional Discount, Total Amount
            #This is the header GRID VIEW Labels
            frame3Label1 = tk.Label(frame3Frame, text = "Bar Code", font = NORMAL_FONT, relief = SUNKEN, width = 15)
            frame3Label1.grid(row = 0, column = 0)

            frame3Label2 = tk.Label(frame3Frame, text = "Product Description", font = NORMAL_FONT, relief = SUNKEN, width = 40)
            frame3Label2.grid(row = 0, column = 1)

            frame3Label3 = tk.Label(frame3Frame, text = "Price", font = NORMAL_FONT, relief = SUNKEN, width = 13)
            frame3Label3.grid(row = 0, column = 2)

            frame3Label4 = tk.Label(frame3Frame, text = "Quantity", font = NORMAL_FONT, relief = SUNKEN, width = 12)
            frame3Label4.grid(row = 0, column = 3)

            
            rowNum = 1
            totalCost = 0
            totalQuantity = 0
            for i in range(len(customerList)):
                if isinstance(customerList[i][4], int):
                    barCodeString = str(customerList[i][4])
                    frame3BarCode = tk.Label(frame3Frame, text = barCodeString, font = NORMAL_FONT, relief = SUNKEN, width = 15)
                    frame3BarCode.grid(row = rowNum, column = 0)

                    prodDesc = customerList[i][1]
                    frame3ProdDesc = tk.Label(frame3Frame, text = prodDesc, font = NORMAL_FONT, relief = SUNKEN, width = 40)
                    frame3ProdDesc.grid(row = rowNum, column = 1)

                    frame3Price = tk.Label(frame3Frame, text = "{:,}".format(customerList[i][2]), font = NORMAL_FONT, relief = SUNKEN, width = 13)
                    frame3Price.grid(row = rowNum, column = 2)

                    totalQuantity += customerList[i][5]
                    frame3Quantity = ttk.Label(frame3Frame, text = "{:,}".format(customerList[i][5]), font = NORMAL_FONT, relief = SUNKEN, width = 12)   #creates an entry box and allows the entry of a string variable
                    frame3Quantity.grid(row = rowNum, column = 3)

                    rowNum += 1        
    #--------------------------------------------------------------------------------------#


    # FRAME4 -------------------------- TOTAL QUANTITY AND AMOUNT ------------------------------------#
            #Fourth frame inside the current MainPage Window Frame
            frame4 = Frame(self, bg = "blue")
            frame4.pack(fill = BOTH)

            frame4Label1 = tk.Label(frame4, text="Total Quantity", font = NORMAL_FONT)
            frame4Label1.grid(row = 0, column = 0, padx = 5, pady = 5)


            frame4Label2 = tk.Label(frame4, text = "{:,}".format(totalQuantity), font = NORMAL_FONT, relief = SUNKEN, width = 16)  #NEED TO GET SUM OF TOTAL QUANTITY
            frame4Label2.grid(row = 0, column = 1)


            #Frame block to separate the first set of labels above from the second set of labels below
            frame4LabelBlock = tk.Label(frame4, width = 16, bg = "blue")
            frame4LabelBlock.grid(row = 0, column = 2, padx = 45, pady = 5)


            frame4Label3 = tk.Label(frame4, text = "Total Cost", font=NORMAL_FONT)
            frame4Label3.grid(row = 0, column = 3, padx = 5, pady = 5)

            frame4Label4 = tk.Label(frame4, text = "N/A", font = NORMAL_FONT, relief = SUNKEN, width = 16)  #NEED TO GET SUM OF TOTAL COST
            frame4Label4.grid(row = 0, column = 4)
    #-------------------------------------------------------------------------------------------------#


    # FRAME5 -------------------------- BAR CODE AND ADD ITEM ------------------------------------#
            #Fifth frame inside the current MainPage Window Frame
            frame5 = Frame(self, bg = "pink")
            frame5.pack(fill = BOTH)

            frame5Label1 = tk.Label(frame5, text = "Product Bar Code", font = NORMAL_FONT)
            frame5Label1.grid(row = 0, column = 0, padx=5, pady=5, sticky = W)

            barCode = tk.StringVar()    #creates the object barCode with a string variable type
            frame5EntryBox = ttk.Entry(frame5, textvariable = barCode, width = 22)   #creates an entry box and allows the entry of a string variable
            frame5EntryBox.focus_set()  #cursor default on entry box
            frame5EntryBox.grid(row = 0, column = 1, padx = 5, pady = 5)

            frame5Spacer = tk.Label(frame5, text = "", font = NORMAL_FONT, bg = "pink")
            frame5Spacer.grid(row = 0, column = 2, padx=20, pady=5, sticky = W)

            frame5Label2 = tk.Label(frame5, text = "Quantity", font = NORMAL_FONT)
            frame5Label2.grid(row = 0, column = 3, padx=5, pady=5, sticky = W)

            quantity = tk.IntVar()    #creates the object barCode with a string variable type
            quantity.set(1)
            frame5EntryBox2 = ttk.Entry(frame5, textvariable = quantity, width = 8)   #creates an entry box and allows the entry of a string variable
            frame5EntryBox2.grid(row = 0, column = 4, padx = 5, pady = 5)

            frame5Button = ttk.Button(frame5, text = "Add Item", command = lambda: updateCustomerList(barCode, quantity))
            frame5Button.grid(row = 0, column = 5, padx = 90, pady = 10)

            self.winfo_toplevel().bind("<Return>", lambda event: updateCustomerList(barCode=barCode, quantity=quantity))    #binds enter/return key to add the barcode given in the entry box 
    #---------------------------------------------------------------------------------------------#


    # FRAME6 -------------------------- PROCESS AND REFRESH ------------------------------------#
            #Sixth frame inside the current MainPage Window Frame
            frame6 = Frame(self, bg = "yellow")
            frame6.pack(fill = BOTH)

            frame6Button = ttk.Button(frame6, text = "Refresh", command=refreshMainFrame)
            frame6Button.grid(row = 0, column = 0, padx = 125, pady = 10)

            frame6Button = ttk.Button(frame6, text = "Process", command=beforeProcessPagePopup)
            frame6Button.grid(row = 0, column = 1, padx = 125, pady = 10)
    #--------------------------------------------------------------------------------------------#   




#error page for when there is a possible error
class ErrorPage(tk.Frame):          #NEED TO CHANGE THIS INTO A POP UP WINDOW INSTEAD OF A FRAME
    def __init__(self, parent, controller):
        tk.Frame.__init__(self,parent)
        
        frame = Frame(self, bg = "green")
        frame.pack(fill = BOTH)

        if errorCode == 1:
            label = tk.Label(frame, text="Error! Please input a correct Master File Document", font=SMALL_FONT)
            label.pack(pady=10, padx=10)

        button1 = ttk.Button(frame, text = "OK", command = lambda: MasterFilePopUp(1)) #One can also do command=quit to quit out of the program
        button1.pack(pady=5, padx=10)




#Start popup to allow the user to either choose Inventory mode or Transaction mode
def startPopup():
    startPopup = tk.Toplevel()
    startPopup.wm_title("FONZY")
    startPopup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

    label = ttk.Label(startPopup, text="Please choose the correct application mode", font=SMALL_FONT)
    label.grid(row = 0, column = 0, pady=10, padx=10, columnspan = 2)

    button1 = ttk.Button(startPopup, text = "Inventory", command = lambda: MasterFilePopUp(0, startPopup) or startPopup.destroy())
    button1.grid(row = 1, column = 0, pady = 10, padx = 10)

    button2 = ttk.Button(startPopup, text = "Transaction", command = lambda: MasterFilePopUp(1, startPopup) or startPopup.destroy())
    button2.grid(row = 1, column = 1, pady = 10, padx = 10)    
    button2.focus_set()  #cursor default on button    



def MasterFilePopUp(mode, startPopup):
    masterPopup = tk.Toplevel()
    masterPopup.wm_title("FONZY")
    masterPopup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

    global modeCode
    modeCode = mode #inherits either 0 or 1
    label = ttk.Label(masterPopup, text = "Please specify the product master file", font = NORMAL_FONT)
    label.pack(side = "top", fill = "x", padx = 5)

    button1 = ttk.Button(masterPopup, text = "Okay", command = lambda: fileExplorer() or masterPopup.destroy() or startPopup.destroy())
    button1.pack(pady = 10)
    button1.focus_set()  #cursor default on button

    if errorCode == 0:
        return 1
    else:
        popupmsg("Please input the correct Master File")
        MasterFilePopUp(mode)




#popup that requires the user to input certain information for receipt generation, default is N/A        
def beforeProcessPagePopup():
    beforeProcessPopup = tk.Toplevel()
    beforeProcessPopup.wm_title("FONZY")
    beforeProcessPopup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

    global tinValue
    global BStyleValue
    global termsValue
    global PWDValue

    label = ttk.Label(beforeProcessPopup, text="TIN", font=SMALL_FONT)
    label.grid(row = 0, column = 0, sticky = "s", pady = 10, padx = 10)

    tinValue = StringVar()
    tinValue.set("N/A")
    entryBox1 = ttk.Entry(beforeProcessPopup, textvariable = tinValue, width = 20)
    entryBox1.grid(row = 0, column = 1, padx = 10, pady = 10)


    label2 = ttk.Label(beforeProcessPopup, text="Business Style", font=SMALL_FONT)
    label2.grid(row = 1, column = 0, sticky = "s", pady = 10, padx = 10)

    BStyleValue = StringVar()
    BStyleValue.set("N/A")
    entryBox2 = ttk.Entry(beforeProcessPopup, textvariable = BStyleValue, width = 20)
    entryBox2.grid(row = 1, column = 1, padx = 10, pady = 10)


    label3 = ttk.Label(beforeProcessPopup, text="Terms", font=SMALL_FONT)
    label3.grid(row = 2, column = 0, sticky = "s", pady = 10, padx = 10)

    termsValue = StringVar()
    termsValue.set("N/A")
    entryBox3 = ttk.Entry(beforeProcessPopup, textvariable = termsValue, width = 20)
    entryBox3.grid(row = 2, column = 1, padx = 10, pady = 10)


    label4 = ttk.Label(beforeProcessPopup, text="OSCA/PWD ID No.", font=SMALL_FONT)
    label4.grid(row = 3, column = 0, sticky = "s", pady = 10, padx = 10)

    PWDValue = StringVar()
    PWDValue.set("N/A")
    entryBox4 = ttk.Entry(beforeProcessPopup, textvariable = PWDValue, width = 20)
    entryBox4.grid(row = 3, column = 1, padx = 10, pady = 10)

    button1 = ttk.Button(beforeProcessPopup, text = "Continue", command = lambda: processPagePopup() or beforeProcessPopup.destroy())
    button1.grid(row = 4, column = 0, padx = 10, pady = 10)

    button2 = ttk.Button(beforeProcessPopup, text = "Exit", command = lambda: beforeProcessPopup.destroy())
    button2.grid(row = 4, column = 1, padx = 10, pady = 10)

    entryBox1.focus_set()  #cursor default on button
    beforeProcessPopup.winfo_toplevel().bind("<Return>", lambda e: processPagePopup() or beforeProcessPopup.destroy())    #binds enter/return key to exit/destroy the popup message


#process purchase page for when a purchase is to be made
def processPagePopup():
    processPopup = tk.Toplevel()
    processPopup.wm_title("FONZY")
    processPopup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

    if modeCode == 1:   #Inventory Mode
        #Need to include check box

        label = ttk.Label(processPopup, text="Payment Type", font=SMALL_FONT)
        label.grid(row = 0, column = 0, sticky = "s", pady = 10, padx = 10)

        #to get exact time, used for invoicing
        time = datetime.datetime.now() #time.time()
        #print(time.time())

        #once button is clicked, it prompts user to find file then it outputs the contents of the file
        #Made all of the checkButtons rely on the checkBox object so it is easier to manage within loops.
        checkBox = ["","","","",""]
        checkBox[0] = StringVar()
        cButton1 = Checkbutton(processPopup, text = "Cash", variable = checkBox[0], onvalue = "CASH", offvalue = "")
        cButton1.grid(row = 1, column = 0, sticky = "W")

        checkBox[1] = StringVar()
        cButton2 = Checkbutton(processPopup, text = "Credit", variable = checkBox[1], onvalue = "CREDIT", offvalue = "")
        cButton2.grid(row = 2, column = 0, sticky = "W")

        checkBox[2] = StringVar()
        cButton3 = Checkbutton(processPopup, text = "Debit", variable = checkBox[2], onvalue = "DEBIT", offvalue = "")
        cButton3.grid(row = 3, column = 0, sticky = "W")

        checkBox[3] = StringVar()
        cButton4 = Checkbutton(processPopup, text = "Check", variable = checkBox[3], onvalue = "CHECK", offvalue = "")
        cButton4.grid(row = 4, column = 0, sticky = "W")

        checkBox[4] = StringVar()
        cButton5 = Checkbutton(processPopup, text = "Salary Deduction", variable = checkBox[4], onvalue = "SALARY DEDUCTION", offvalue = "")
        cButton5.grid(row = 5, column = 0, rowspan = 2)


        label2 = ttk.Label(processPopup, text="Customer Type", font=SMALL_FONT)
        label2.grid(row = 7, column = 0, sticky = "s", pady = 10, padx = 10)

        #Determines customer type, default is new
        listBox = Listbox(processPopup, selectmode = SINGLE, width = 0, height = 5)
        listBox.insert(1, "New")
        listBox.insert(2, "Old")
        listBox.insert(3, "Employee")
        listBox.insert(4, "Guest")
        listBox.insert(5, "Non-stat/Intern")
        listBox.grid(row = 8, column = 0, sticky = "nsew", padx = 10)

        button = ttk.Button(processPopup, text = "OK", command = lambda: printString() or processPopup.destroy())
        button.grid(row = 10, column = 0, pady = 20, padx = 10)

        button.focus_set()  #cursor default on button
        processPopup.winfo_toplevel().bind("<Return>", lambda e: printString() or processPopup.destroy())    #binds enter/return key to exit/destroy the popup message

        def printString():
            global paymentString
            global customerType
            customerType = listBox.get(ACTIVE)
            for i in range(len(checkBox)):  #sets payment string to the types of payments going to be used
                if checkBox[i].get() != "":
                    paymentString += checkBox[i].get()
                    paymentString += "  "
            
            finalPayment(checkBox)
            return

    else:   #Transaction mode
        #UPDATE EXCEL FILE then ASK USER ONE MORE TIME BEFORE EXIT
        label = ttk.Label(processPopup, text="Are you sure you are finished taking inventory?", font=SMALL_FONT)
        label.grid(row = 0, column = 0, pady=10, padx=10, columnspan = 2)

        button1 = ttk.Button(processPopup, text = "Yes", command = lambda: updateExcel() or processPopup.destroy())
        button1.grid(row = 1, column = 0, pady = 10, padx = 10)

        button2 = ttk.Button(processPopup, text = "No", command = lambda: refreshMainFrame() or processPopup.destroy())
        button2.grid(row = 1, column = 1, pady = 10, padx = 10)

        button1.focus_set()  #cursor default on button
        processPopup.winfo_toplevel().bind("<Return>", lambda e: updateExcel() or processPopup.destroy())    #binds enter/return key to exit/destroy the popup message

        def updateExcel():
            
            wbName = filename.split("/")    #gets masterFile name
            wbNameString = wbName[-1]   #gets last element in list

            wb=load_workbook(wbNameString)  #opens masterFile
            activeWS=wb.active   #uses active workbook for edits

            for r in range(0,len(customerList)):
                activeWS.cell(row=r+1,column=6).value=customerList[r][5]

            wb.save(wbNameString)  #saves masterFile with edits
            return



#Payment page after Payment type popup
def finalPayment(checkBox):
    paymentPopup = tk.Toplevel()
    paymentPopup.wm_title("FONZY")
    paymentPopup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

    costTotal = 0
    for i in range(len(customerList)):  #finds total price of transaction
        costTotal += (customerList[i][6]*customerList[i][2])-((customerList[i][3]/100)*(customerList[i][6]*customerList[i][2]))


    #variable declarations
    cashAmount = IntVar()
    cashAmount.set(0)
    creditAmount = IntVar()
    creditAmount.set(0)
    debitAmount = IntVar()
    debitAmount.set(0)
    checkAmount = IntVar()
    checkAmount.set(0)
    sDeductionAmount = IntVar()
    sDeductionAmount.set(0)

    def finalPaymentBuild(checkBox, stateBuild):

        label = ttk.Label(paymentPopup, text="Total Amount", font=SMALL_FONT)
        label.grid(row = 0, column = 0, pady=10, padx=10, sticky = "W")

        label2 = ttk.Label(paymentPopup, text="{:,}".format(costTotal), font=SMALL_FONT)
        label2.grid(row = 0, column = 1, pady=10, padx=10, sticky = "W", columnspan = 2)

        #CASH
        label3 = ttk.Label(paymentPopup, text="Cash", font=SMALL_FONT)
        label3.grid(row = 1, column = 0, pady=10, padx=10, sticky = "W")    

        EntryBox1 = ttk.Entry(paymentPopup, textvariable = cashAmount, width = 15)
        if checkBox[0].get() == "":    #disables entry box if checkbox isn't == 1
            EntryBox1.config(state="disabled")
        EntryBox1.grid(row = 1, column = 1, padx = 5, pady = 5, sticky = "W", columnspan = 2)


        #CREDIT
        label4 = ttk.Label(paymentPopup, text="Credit", font=SMALL_FONT)
        label4.grid(row = 2, column = 0, pady=10, padx=10, sticky = "W")    

        EntryBox2 = ttk.Entry(paymentPopup, textvariable = creditAmount, width = 15)
        if checkBox[1].get() == "":    #disables entry box if checkbox isn't == 1
            EntryBox2.config(state="disabled")    
        EntryBox2.grid(row = 2, column = 1, padx = 5, pady = 5, sticky = "W", columnspan = 2)


        #DEBIT
        label5 = ttk.Label(paymentPopup, text="Debit", font=SMALL_FONT)
        label5.grid(row = 3, column = 0, pady=10, padx=10, sticky = "W")    

        EntryBox3 = ttk.Entry(paymentPopup, textvariable = debitAmount, width = 15)
        if checkBox[2].get() == "":    #disables entry box if checkbox isn't == 1
            EntryBox3.config(state="disabled")
        EntryBox3.grid(row = 3, column = 1, padx = 5, pady = 5, sticky = "W", columnspan = 2)


        #CHECK
        label6 = ttk.Label(paymentPopup, text="Check", font=SMALL_FONT)
        label6.grid(row = 4, column = 0, pady=10, padx=10, sticky = "W")    

        EntryBox4 = ttk.Entry(paymentPopup, textvariable = checkAmount, width = 15)
        if checkBox[3].get() == "":    #disables entry box if checkbox isn't == 1
            EntryBox4.config(state="disabled")
        EntryBox4.grid(row = 4, column = 1, padx = 5, pady = 5, sticky = "W", columnspan = 2)


        #SALARY DEDUCTION
        label7 = ttk.Label(paymentPopup, text="Salary Deduction", font=SMALL_FONT)
        label7.grid(row = 5, column = 0, pady=10, padx=10, sticky = "W")    

        EntryBox5 = ttk.Entry(paymentPopup, textvariable = sDeductionAmount, width = 15)
        if checkBox[4].get() == "":    #disables entry box if checkbox isn't == 1
            EntryBox5.config(state="disabled")
        EntryBox5.grid(row = 5, column = 1, padx = 5, pady = 5, sticky = "W", columnspan = 2)


        #TOTAL CUSTOMER PAYMENT (only updated after Calculate button is pressed)
        label8 = ttk.Label(paymentPopup, text="Total Customer Payment", font=SMALL_FONT)
        label8.grid(row = 6, column = 0, pady=10, padx=10, sticky = "W")

        totalCP = IntVar()
        totalCP = 0
        if stateBuild == 1:
            global totalCustomerPayment
            totalCP = cashAmount.get() + debitAmount.get() + creditAmount.get() + checkAmount.get() + sDeductionAmount.get()    #gets sum of list
            totalCustomerPayment = totalCP

        label9 = ttk.Label(paymentPopup, text="{:,}".format(totalCP), font=SMALL_FONT)
        label9.grid(row = 6, column = 1, pady=10, padx=10, sticky = "W", columnspan = 2)        


        #CHANGE
        label8 = ttk.Label(paymentPopup, text="Change", font=SMALL_FONT)
        label8.grid(row = 7, column = 0, pady=10, padx=10, sticky = "W")

        totalChange = IntVar()
        totalChange = 0
        if stateBuild == 1: #calculates totalChange
            totalChange = totalCP - costTotal    #gets total change (customer payment - total cost)

        label9 = ttk.Label(paymentPopup, text="{:,}".format(totalChange), font=SMALL_FONT)
        label9.grid(row = 7, column = 1, pady=10, padx=10, sticky = "W", columnspan = 2)

        #sets variable states back to 0 to refresh data input
        state = 0

        #BUTTONS Calculate, OK, and Cancel
        button1 = ttk.Button(paymentPopup, text = "Calculate", command = lambda: finalPaymentBuild(checkBox, 1))
        button1.grid(row = 8, column = 0, pady = 10, padx = 10)    

        button2 = ttk.Button(paymentPopup, text = "Proceed", command = lambda: paymentContinue() or paymentPopup.destroy())
        button2.grid(row = 8, column = 1, pady = 10, padx = 10)

        button3 = ttk.Button(paymentPopup, text = "Cancel", command = lambda: paymentPopup.destroy())
        button3.grid(row = 8, column = 2, pady = 10, padx = 10)

        button1.focus_set()  #cursor default on button
        paymentPopup.winfo_toplevel().bind("<Return>", lambda e: finalPaymentBuild(checkBox, 1))    #binds enter/return key to exit/destroy the popup message

    finalPaymentBuild(checkBox, 0)  #starts payment calculator




#Continue question after Payment popup
def paymentContinue():
    pContinuePopup = tk.Toplevel()
    pContinuePopup.wm_title("FONZY")
    pContinuePopup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

    label = ttk.Label(pContinuePopup, text="Are you sure about the purchase details?", font=SMALL_FONT)
    label.grid(row = 0, column = 0, pady=10, padx=10, columnspan = 2)

    button1 = ttk.Button(pContinuePopup, text = "Yes", command = lambda: excelCheckoutUpdate() or pContinuePopup.destroy())
    button1.grid(row = 1, column = 0, pady = 10, padx = 10)    

    button2 = ttk.Button(pContinuePopup, text = "No", command = lambda: pContinuePopup.destroy())
    button2.grid(row = 1, column = 1, pady = 10, padx = 10)    

    button1.focus_set()  #cursor default on button
    pContinuePopup.winfo_toplevel().bind("<Return>", lambda e: excelCheckoutUpdate() or pContinuePopup.destroy())    #binds enter/return key to exit/destroy the popup message

    now = datetime.datetime.now()   #gets date-time
    nowDate = now.strftime("%d-%m-%Y")
    nowTime = now.strftime("%H:%M")
    nowTimeWord = now.strftime("%H;%M")
    #Updates excel file to accomodate customer purchases
    def excelCheckoutUpdate():
        #Gets last row of workbook to figure out if program needs to create headers (max row <= 1) or just add to current workbook
        #opens and edits workbook
        excelFilePathArray = filename.split("/")
        excelFilePathArray = excelFilePathArray[:-1]
        excelFilePath = '/'.join(excelFilePathArray)
        excelFile = excelFilePath + "/" + excelString

        wb = load_workbook(excelFile)
        ws = wb.active
        
        row_count = ws.max_row + 1
        col_count = 1

        if row_count <= 2:  #headers: Customer Name | Barcode Number | Product Description | Price | Quantity | Total Amount | Customer Type | Payment Form | Date | Time
            ws.cell(row = 1, column = col_count).value = "Customer Name"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Barcode Number"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Product Description"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Price"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Quantity"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Total Amount"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Customer Type"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Payment Form"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Date"
            col_count += 1
            ws.cell(row = 1, column = col_count).value = "Time"

            col_count = 1  #resets column back to zero

        #inputs customer purchase data
        for r in range(0,len(customerList)):
            ws.cell(row = row_count, column = col_count).value = customerName.get() #customer name
            col_count += 1
            ws.cell(row = row_count, column = col_count).value = customerList[r][0] #barcode number
            col_count += 1
            ws.cell(row = row_count, column = col_count).value = customerList[r][1] #product description
            col_count += 1
            ws.cell(row = row_count, column = col_count).value = customerList[r][2] #product price
            col_count += 1
            ws.cell(row = row_count, column = col_count).value = customerList[r][6] #product quantity
            col_count += 1
            totAmount = customerList[r][2] * customerList[r][6]
            ws.cell(row = row_count, column = col_count).value = totAmount  #product total cost
            col_count += 1
            ws.cell(row = row_count, column = col_count).value = customerType    #customer type
            col_count += 1
            ws.cell(row = row_count, column = col_count).value = paymentString
            col_count += 1
            ws.cell(row = row_count, column = col_count).value = nowDate
            col_count += 1
            ws.cell(row = row_count, column = col_count).value = nowTime
            
            row_count += 1  #updates row_count
            col_count = 1  #resets column back to zero

        wb.save(excelFile)
        wordCheckoutUpdate()    #Creates printable word document receipt

    def wordCheckoutUpdate():
        wordString = customerName.get() + "," + "date=" + str(nowDate) + "," + "time=" +str(nowTimeWord) + ".docx"

        wordFilePathArray = filename.split("/")
        wordFilePathArray = wordFilePathArray[:-1]
        wordFilePath = '/'.join(wordFilePathArray)
        wordFile = wordFilePath + "/" + wordString

        #adds paragraph spacing
        def addParagraphSpace(numberOfSpaces):
            #Adds 9 blank paragraphs for structure
            for i in range(numberOfSpaces):
                blankParagraph = document.add_paragraph()
                blankParagraph.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
                pFormat = blankParagraph.paragraph_format
                pFormat.space_before = Pt(0)
                pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
            return

        document = Document()

        #Sets the font and size so that it will be monospace
        style = document.styles["Normal"]
        font = style.font
        font.name = "Consolas"
        font.size = Pt(10)

        addParagraphSpace(5)    #Adds 8 blank paragraphs for structure

        #THERER ARE 78 charaters per line
        #52 spaces before top right section
        #57 for date, 58 for terms, 69 for OSCA,ID number
        #22 paragraphs for items only! i.e. max 20 items

        #Cashier  name
        cashierStringName = "Cashier: " + str(cashierString.get())
        pCashierString = buildCharacterParagraphArray(0, cashierStringName, 76, 0, "")
        pCashierName = document.add_paragraph()
        pCashierName.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pCashierName.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pCashierName.add_run(pCashierString)

        #Time
        timeString = "Time: " + str(nowTime)
        pTimeString = buildCharacterParagraphArray(0, timeString, 76, 0, "")
        pTime = document.add_paragraph()
        pTime.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pTime.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pTime.add_run(pTimeString)

        #Phone number
        phoneString = "Customer Phone: " + str(phone.get())
        pPhone = buildCharacterParagraphArray(0, phoneString, 76, 0, "")
        pPhoneNumber = document.add_paragraph()
        pPhoneNumber.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pPhoneNumber.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pPhoneNumber.add_run(pPhone)

        #1ST PART OF THE LINE IS 50 SPACES, THEN 2ND PART OF THE LINE IS 27 SPACES
        #8 spaces before name input; max 42 character for name; 8 spaces before date input
        #Name and date paragraph
        p1 = buildCharacterParagraphArray(8, list(str(customerName.get())), 42, 8, list(str(nowDate)))
        pNameDate = document.add_paragraph()
        pNameDate.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pNameDate.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pNameDate.add_run(p1)

        #9 spaces for address; max 41 characters for address; 10 spaces before terms input
        p2 = buildCharacterParagraphArray(9, list(str(address.get())), 41, 8, str(termsValue.get()))
        pAddressTerms = document.add_paragraph()
        pAddressTerms.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pAddressTerms.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pAddressTerms.add_run(p2)                

        #5 spaces for TIN; max 45 characters for TIN; 20 spaces before PWD ID input; max 8 spaces before PWD ID input
        p3 = buildCharacterParagraphArray(5, str(tinValue.get()), 45, 20, str(PWDValue.get()))
        pTinPwd = document.add_paragraph()
        pTinPwd.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pTinPwd.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pTinPwd.add_run(p3)

        #13 spaces for TIN; max 37 characters for TIN; 20 spaces before PWD ID input; max 8 spaces before PWD ID input
        p4 = buildCharacterParagraphArray(13, str(BStyleValue.get()), 37, 22,"")
        pBStyleSignature = document.add_paragraph()
        pBStyleSignature.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pBStyleSignature.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pBStyleSignature.add_run(p4)

        addParagraphSpace(1)    #Adds 2 blank paragraphs for structure

        #TITLE INPUT
        pTitle = document.add_paragraph()
        pTitle.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pTitle.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pTitle.add_run("Item Number       Description                   Qty    Price      Amount")

        addParagraphSpace(1)    #Adds 1 blank paragraphs for structure

        #CREATES BODY INPUT PARAGRAPHS
        for i in range(len(customerList)):
            pBody = document.add_paragraph()
            pBody.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
            pFormat = pBody.paragraph_format
            pFormat.space_before = Pt(0)
            pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
            customerBarCodeString = str(customerList[i][4]).strip()    #strips the EAN bar code string value to make sure that there aren't any white spaces
            customerDescriptionString = str(customerList[i][1]).strip()    #strips the Material Description string value to make sure that there aren't any white spaces
            pBodyArrayList = buildBodyParagraphArray(i, customerBarCodeString, customerDescriptionString)
            pBody.add_run(pBodyArrayList)


        quantityTotal = 0
        costTotal = 0
        for a in range(len(customerList)):  #finds total price of transaction
            quantityTotal += customerList[a][6]
            costTotal += (customerList[a][6]*customerList[a][2])-((customerList[a][3]/100)*(customerList[a][6]*customerList[a][2]))            


        #converts quantity and cost to string, ready for char input
        totalQuantityString = str(quantityTotal)
        totalCostString = str(costTotal)
        
        #paragraph for total number of items sold and price, difference is the 5PT space before and after
        pTotal = document.add_paragraph()
        pTotal.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pTotal.paragraph_format
        pFormat.space_before = Pt(5)
        pFormat.space_after = Pt(5)    #sets line spacing to 0 instead of the default 1.5
        pTotal.add_run("                                        Total   ")

        #adds total quantity to the single line paragraph; has 18 allotted spaces
        for k in range(len(totalQuantityString)):
            pTotal.add_run(totalQuantityString[k])
        if (18 - len(totalQuantityString)) > 0: 
            for l in range(18 - len(totalQuantityString)):  #adds space in the paragraph
                pTotal.add_run(" ")
        #adds total amount/price to the single line paragraph; has 12 allotted spaces
        for m in range(len(totalCostString)):
            pTotal.add_run(totalCostString[m])


        #adds spaces after total quantity and amount is inputted into a paragraph
        if (21 - len(customerList)) > 0:
            spaceAfterCustomerList = 21 - len(customerList)
            addParagraphSpace(spaceAfterCustomerList)   #adds a dynamically sized blank paragraphs for structure


        #TAX INFO
        #VATable Sales = costTotal / 1.12
        #Add: Vat Amount = costTotal - VATable Sales
        #Total Amount due = costTotal
        VATableSales = costTotal / 1.12 #gets total price for item without VAT
        VATableSales = round(VATableSales, 2)
        VATAmount = costTotal - VATableSales    #gets VAT price for item
        VATAmount = round(VATAmount, 2)

        #68 spaces for VATable sales; max 12 characters for VATable sales
        pVAT = buildCharacterParagraphArray(66, str(VATableSales), 12, 0, "")
        pVATableSales = document.add_paragraph()
        pVATableSales.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pVATableSales.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pVATableSales.add_run(pVAT)

        #paragraph space
        blankParagraph = document.add_paragraph()
        blankParagraph.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = blankParagraph.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        
        #paragraph space
        blankParagraph = document.add_paragraph()
        blankParagraph.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = blankParagraph.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5

        #64 spaces for VATable sales; max 12 characters
        pTSales = buildCharacterParagraphArray(64, str(costTotal), 14, 0, "")
        pTotalSales = document.add_paragraph()
        pTotalSales.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pTotalSales.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pTotalSales.add_run(pTSales)

        #64 spaces for VATable sales; max 12 characters
        pVAmount = buildCharacterParagraphArray(68, str(VATAmount), 10, 0, "")
        pVATAmount = document.add_paragraph()
        pVATAmount.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pVATAmount.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pVATAmount.add_run(pVAmount)

        #paragraph space for amount paid, difference is the 5PT space after
        blankParagraph = document.add_paragraph()
        blankParagraph.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = blankParagraph.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(5)    #sets line spacing to 5 instead of the default 1.5        

        #68 spaces for VATable sales; max 12 characters
        pADue = buildCharacterParagraphArray(56, str(round(costTotal, 2)), 22, 0, "")
        pAmountDue = document.add_paragraph()
        pAmountDue.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pAmountDue.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pAmountDue.add_run(pADue)


        #65 spaces for Customer Payment; max 13 characters
        pTotalCP = buildCharacterParagraphArray(65, str(round(totalCustomerPayment, 2)), 13, 0, "")
        pTotalCustomerPayment = document.add_paragraph()
        pTotalCustomerPayment.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pTotalCustomerPayment.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pTotalCustomerPayment.add_run(pTotalCP)

        #65 spaces for Customer Change; max 13 characters
        totalCustomerChange = totalCustomerPayment - costTotal
        pTotalChange = buildCharacterParagraphArray(65, str(round(totalCustomerChange, 2)), 13, 0, "")
        pTotalCustomerChange = document.add_paragraph()
        pTotalCustomerChange.style = document.styles["Normal"]    #sets the style to match the given monospace style mentioned above
        pFormat = pTotalCustomerChange.paragraph_format
        pFormat.space_before = Pt(0)
        pFormat.space_after = Pt(0)    #sets line spacing to 0 instead of the default 1.5
        pTotalCustomerChange.add_run(pTotalChange)        

        #saves document
        document.save(wordFile)
        printCheckoutUpdate(wordFile)   #allows user to print document




#allows the user to print or reprint the current receipt
def printCheckoutUpdate(wordFile):
    #NEED TO WATCH VIDEO ABOUT PRINTING WORD DOCUMENT!
    clearCustomerInfo()




#builds a one line generalized paragraph
def buildBodyParagraphArray(customerListIndexValue, customerBarCodeString, customerDescriptionString):
    arrayList = []
    #APPENDS BARCODE
    for i in range(15):
        if i < len(customerBarCodeString):
            arrayList.append(customerBarCodeString[i])
        else:
            arrayList.append(" ")
    for j in range(3):
        arrayList.append(" ")

    #APPENDS ITEM DESCRIPTION STIRNG
    for k in range(27):
        if k < len(customerDescriptionString):
            arrayList.append(customerDescriptionString[k])
        else:
            arrayList.append(" ")
    for l in range(3):
        arrayList.append(" ")    
        
    #APPENDS QUANTITY
    quantityString = str(customerList[customerListIndexValue][6])
    for m in range(4):
        if m < len(quantityString):
            arrayList.append(quantityString[m])
        else:
            arrayList.append(" ")
    for n in range(3):
        arrayList.append(" ")

    #APPENDS PRICE
    priceString = str(customerList[customerListIndexValue][2])
    #priceString = "{:,}".format(priceStringValue)
    for p in range(8):
        if p < len(priceString):
            arrayList.append(priceString[p])
        else:
            arrayList.append(" ")
    for q in range(3):
        arrayList.append(" ")

    #APPENDS AMOUNT
    amount = (customerList[customerListIndexValue][6]*customerList[customerListIndexValue][2])-((customerList[customerListIndexValue][3]/100)*(customerList[customerListIndexValue][6]*customerList[customerListIndexValue][2]))
    amountString = str(amount)
    #amountString = "{:,}".format(amountStringValue)
    for r in range(12):
        if r < len(amountString):
            arrayList.append(amountString[r])
        else:
            pass

    return arrayList




#builds a one line generalized paragraph
def buildCharacterParagraphArray(initialSpaceNumber, initialStringValue, stringSpaceAllowanceNumber, secondarySpaceNumber, finalStringValue):
    #builds array with 78 character elements
    j = 0
    k = 0
    arrayList = []

    for i in range(78):
        if i <= initialSpaceNumber:
            arrayList.append(" ")
        
        else:
            if i < (stringSpaceAllowanceNumber+initialSpaceNumber):
                if j < len(initialStringValue):
                    arrayList.append(initialStringValue[j])
                    j += 1
                else:
                    arrayList.append(" ")

            else:
                if i < (stringSpaceAllowanceNumber + initialSpaceNumber + secondarySpaceNumber):
                    arrayList.append(" ")
                
                else:
                    if k < len(finalStringValue):
                        arrayList.append(finalStringValue[k])
                        k += 1
                    else:
                        arrayList.append(" ")
    return arrayList




#function that clears customer info for next purchase
def clearCustomerInfo():
    global customerList
    global customerName
    global phone
    global address
    global customerType

    #resets customer info
    customerList[:] = []    #deletes all elements in list and recreates the list
    
    customerName = StringVar()
    customerName.set("")
    phone = IntVar()
    phone.set(0)
    address = StringVar()
    address.set("")

    customerType = "New"

    refreshMainFrame()
    return




#opens file explorer for the cashier to choose which masterfile to use
def fileExplorer():
    #intializes another instance of tkinter
    global masterList   #allows user to add item to masterList global variable
    global customerList
    global filename
    try:    
        filename = filedialog.askopenfilename(initialdir = "/",title = "Select file",filetypes = (("XLSX files","*.xlsx"),("All Files","*.*")))
            
        if re.match("[A-Za-z0-9]",filename):
            sheet = pe.get_array(file_name=filename) #puts data into array
            masterList = sheet
            errorCode = 0 #resets errorCode
            if modeCode == 1:
                cashierStartPopup()

            if modeCode == 0:   #sets customer list dependent on master list data with row number on last index element of customer list
                for i in range(len(masterList)):    #sets customerList to equal values of masterList
                    customerList.append(masterList[i][:])
                    customerList[i].append(i+1)

                refreshMainFrame()
        else:
            errorCode = 1
    
    except ValueError:
        errorCode = 1




#Creates popup frame to get Cashier name, Event name and date
def cashierStartPopup():
    cashierPopup = tk.Toplevel()
    cashierPopup.wm_title("FONZY")
    cashierPopup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

    global cashierString

    #For cashier name
    label = ttk.Label(cashierPopup, text = "Cashier Name: ", font = NORMAL_FONT)
    label.grid(row = 0, column = 0, sticky = "nsew", padx = 5, pady = 5)

    cashierString = tk.StringVar()
    cashierString.set("Cashier Name")
    cashierEntryBox = ttk.Entry(cashierPopup, textvariable = cashierString, width = 25)
    cashierEntryBox.focus_set() #sets cursor default to this box
    cashierEntryBox.grid(row = 0, column = 1, padx = 5, pady = 5)

    #For event name
    label2 = ttk.Label(cashierPopup, text = "Event Name:  ", font = NORMAL_FONT)
    label2.grid(row = 1, column = 0, sticky = "nsew", padx = 5, pady = 5)

    eventString = tk.StringVar()
    eventString.set("Event Name")
    eventEntryBox = ttk.Entry(cashierPopup, textvariable = eventString, width = 25)
    eventEntryBox.grid(row = 1, column = 1, padx = 5, pady = 5)

    #To exit program, goes to masterFilePopup
    button1 = ttk.Button(cashierPopup, text = "Okay", command = lambda: setProgramStartData() or cashierPopup.destroy())
    button1.grid(row = 2, padx = 5, pady = 5, sticky = "nsew")

    cashierEntryBox.focus_set()  #cursor default on button
    cashierPopup.winfo_toplevel().bind("<Return>", lambda e: setProgramStartData() or cashierPopup.destroy())    #binds enter/return key to exit/destroy the popup message

    #sets the data taken from the entry boes and sets it into the appropriate objects
    def setProgramStartData():
        global excelString
        #gets excel string to create excel file for the day
        excelString = cashierString.get() + "-" + eventString.get() + "-" + str(datetime.datetime.today().strftime('%d,%m,%Y') + ".xlsx")
        
        #creates work book inside the filepath given, need to fix filename to imitate the same filepath of the materfilePopup using .split()
        excelFilePathArray = filename.split("/")
        excelFilePathArray = excelFilePathArray[:-1]
        excelFilePath = '/'.join(excelFilePathArray)
        #checks if the file already exists, if it doesn't then it creates the file, if it does then it continues on in the program.
        filePath = Path(excelFilePath + "/" + excelString)

        if not filePath.is_file():
            workbook = xw.Workbook(excelFilePath + "/" + excelString)
            worksheet = workbook.add_worksheet()

            workbook.close()




#Creates popup message bars
def popupmsg(msg):
    popup = tk.Toplevel()
    popup.wm_title("FONZY")
    popup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage
    
    label = ttk.Label(popup, text = msg, font = NORMAL_FONT)
    label.pack(side = "top", fill = "x", pady = 10)

    button1 = ttk.Button(popup, text = "Okay", command = popup.destroy)
    button1.pack()

    button1.focus_set()  #cursor default on button
    denyPopup.winfo_toplevel().bind("<Return>", lambda e: popup.destroy())    #binds enter/return key to exit/destroy the popup message




def updateCustomerList(barCode, quantity):
    global customerList   #global variable to allow user to update updateCustomerList

    if modeCode == 1:   #Transaction Mode
    
        if (len(customerList) == 20) and (quantity.get() > 0): #cashier is notified that max number of items is already on the list
            denyUpdateCustomerList()
        else:    
            if not barCode.get():   #checks if barCode is empty
                return  
            else:   #finds barCode inside masterList
                for i in range(len(masterList)):    #searches through master list to see if barCode is inside masterList
                   
                    if int(barCode.get()) == masterList[i][4]:   #if bar code is inside the masterList

                        if len(customerList) == 0:   #if customerList is empty
                            customerList.append(masterList[i][:])  #adds a masterList object inside customerList
                            customerList[0].append(quantity.get())   #gives a quantifiable value to number of products the customer wants to purchase
                            
                            if customerList[0][6] <= 0: #deletes element if item quantity value is 0 or less than 0
                                del customerList[0]
                            refreshMainFrame()  #sends back to MainPage Frame

                        else:   #if customerList is not empty
                            for j in range(len(customerList)):    #searches through customerList to see if item is already inside; checks for repeats
                                
                                if int(barCode.get()) == customerList[j][4]:   #if is a repeated barCode
                                    customerList[j][6] += quantity.get()
                                    
                                    if customerList[j][6] <= 0: #deletes element if item quantity value is 0 or less than 0
                                        del customerList[j]
                                    refreshMainFrame()  #sends back to MainPage Frame
                                    return

                            customerList.append(masterList[i][:])  #adds a masterList object inside customerList
                            customerList[len(customerList)-1].append(quantity.get())   #gives a quantifiable value to number of products the customer wants to purchase
                            
                            if customerList[len(customerList)-1][6] <= 0: #deletes element if item quantity value is 0 or less than 0
                                del customerList[len(customerList)-1]
                            refreshMainFrame()  #sends back to MainPage Frame
    else:   #Inventory Mode
        if not barCode.get():   #checks if barCode is empty
            return  
        
        else:   #finds barCode inside customerList

            for i in range(len(customerList)):    #searches through master list to see if barCode is inside customerList

                if int(barCode.get()) == customerList[i][4]:   #if bar code is inside the customerList    
                    customerList[i][5] += quantity.get()

                    if customerList[i][5] <= 0: #sets quantity field back to 0 if the value is a negative number, doesn't delete the field entirely
                        customerList[i][5] = 0
                    
                    refreshMainFrame()  #sends back to MainPage Frame
                    return

    refreshMainFrame()
    return            



#user is notified that he/she cannot add more items inside the list
def denyUpdateCustomerList():
    denyPopup = tk.Toplevel()
    denyPopup.wm_title("FONZY")
    denyPopup.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

    label = ttk.Label(denyPopup, text = "Customer purchase list has reached maximum capacity", font = NORMAL_FONT)
    label.pack(side = "top", padx = 10)

    label2 = ttk.Label(denyPopup, text = "Please delete some items or process this purchase", font = NORMAL_FONT)
    label2.pack(side = "top", padx = 10)

    button = ttk.Button(denyPopup, text = "Okay", command = denyPopup.destroy)
    button.pack()    

    button.focus_set()  #cursor default on button
    denyPopup.winfo_toplevel().bind("<Return>", lambda e: denyPopup.destroy())    #binds enter/return key to exit/destroy the popup message



#deletes and rebuilds the app i.e. refresh
def refreshMainFrame():
    global app
    app.frames[MainPage].destroy()
    app.frames[MainPage] = MainPage(container, app)
    app.frames[MainPage].grid(row=0, column = 0, sticky = "nsew")
    app.frames[MainPage].tkraise()





#-------------------------------- START OF PROGRAM ----------------------------------------#

#runs tkinter program
app = POS()
#"width x height"
app.geometry("700x700") #makes app into a 700x700p screen, can change size to liking
app.resizable(False, False) #window isn't resizable. Makes it easier for the owner to manage

#finds Master File first
# run after `POS` will be created and `app.mainloop()` will start
app.after(100, startPopup)
app.mainloop()